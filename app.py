import streamlit as st
from pdf2docx import Converter
from pdf2image import convert_from_bytes
from docx import Document
from docx.shared import Inches
from io import BytesIO
import tempfile
import os
import time

st.set_page_config(page_title="PDF → Word Converter", layout="centered")

st.title("📄 ➜ 📝 PDF → Word (2 chế độ: Visual / Text)")
st.markdown("Chọn chế độ **Visual** để giữ nguyên công thức dưới dạng ảnh (khuyến nghị cho PDF có công thức). Chọn **Text** để xuất text (có thể mất công thức/phức tạp).")

mode = st.radio("Chọn chế độ chuyển đổi:", ("Visual (ảnh, giữ nguyên bố cục)", "Text (pdf2docx)"))

uploaded_file = st.file_uploader("Tải lên file PDF", type=["pdf"])
quality = st.slider("Độ phân giải ảnh khi dùng chế độ Visual (dpi)", min_value=100, max_value=400, value=200)
max_pages = st.number_input("Số trang tối đa (0 = tất cả)", min_value=0, value=0, step=1)

if uploaded_file is not None:
    # Hiển thị thông tin file
    file_size_kb = len(uploaded_file.getvalue()) // 1024
    st.write(f"**Tên file:** {uploaded_file.name} — **Kích thước:** {file_size_kb} KB")

    if st.button("Bắt đầu chuyển đổi"):
        start_time = time.time()
        fname = os.path.splitext(uploaded_file.name)[0]
        if mode.startswith("Visual"):
            st.info("Chế độ Visual: sẽ render từng trang thành ảnh rồi chèn vào .docx")
            try:
                with st.spinner("Đang render PDF thành ảnh... (cần poppler trên hệ thống)"):
                    # convert_from_bytes returns PIL Image list
                    # respect max_pages if >0
                    pdf_bytes = uploaded_file.getvalue()
                    images = convert_from_bytes(pdf_bytes, dpi=quality)
                    if max_pages and max_pages > 0:
                        images = images[:max_pages]
                st.success(f"Đã render {len(images)} trang thành ảnh.")
            except Exception as e:
                st.error(f"Lỗi khi render PDF: {e}")
                raise

            # Tạo docx và chèn từng ảnh
            doc = Document()
            # Optional: set page margins or style if needed
            tmp_docx = BytesIO()
            progress = st.progress(0)
            for i, img in enumerate(images, start=1):
                # save image to BytesIO as PNG to then insert
                img_byte = BytesIO()
                img.save(img_byte, format="PNG")
                img_byte.seek(0)

                # Add a page break before second+ pages to keep pages separate
                if i > 1:
                    doc.add_page_break()

                # Insert image: adjust width to page width approx (use Inches)
                # typical Word page width minus margins ~ 6.5 inches -> adjust to fit
                try:
                    doc.add_picture(img_byte, width=Inches(6.8))
                except Exception:
                    # fallback without width
                    doc.add_picture(img_byte)

                progress.progress(int(i / len(images) * 100))

            # Save to bytes
            doc_stream = BytesIO()
            doc.save(doc_stream)
            doc_stream.seek(0)

            st.success("Hoàn thành chuyển đổi (Visual → .docx).")
            elapsed = time.time() - start_time
            st.write(f"Thời gian: {elapsed:.1f}s")

            st.download_button(
                label="Tải về file .docx (Visual)",
                data=doc_stream,
                file_name=f"{fname}_visual.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        else:
            # Text mode: use pdf2docx
            st.info("Chế độ Text: dùng pdf2docx để convert sang docx (có thể không giữ chính xác công thức).")
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    pdf_path = os.path.join(tmpdir, "input.pdf")
                    out_path = os.path.join(tmpdir, "output.docx")
                    # write uploaded bytes to file
                    with open(pdf_path, "wb") as f:
                        f.write(uploaded_file.getvalue())

                    converter = Converter(pdf_path)
                    # if user wants limit pages
                    if max_pages and max_pages > 0:
                        converter.convert(out_path, start=0, end=max_pages-1)
                    else:
                        converter.convert(out_path)
                    converter.close()

                    # read output
                    with open(out_path, "rb") as f:
                        docx_bytes = f.read()

                st.success("Hoàn thành chuyển đổi (Text → .docx).")
                elapsed = time.time() - start_time
                st.write(f"Thời gian: {elapsed:.1f}s")

                st.download_button(
                    label="Tải về file .docx (Text)",
                    data=docx_bytes,
                    file_name=f"{fname}_text.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"Lỗi khi convert bằng pdf2docx: {e}")
